import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsmap
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
from xgboost import XGBClassifier
from lightgbm import LGBMClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, Normalizer
from sklearn.ensemble import RandomForestClassifier, AdaBoostClassifier
from sklearn.metrics import classification_report, accuracy_score, ConfusionMatrixDisplay, confusion_matrix

# Load the dataset
title = "4 model on BreastCancerWisconsin(Diagnostic) SMOTETomek"
doc = Document()
doc.add_heading(title, level=1)

df_src = "SMOTETomek-report/SMOTETomek-BreastCancerWisconsin(Diagnostic).csv"
df = pd.read_csv(df_src)

X = df.drop(["diagnosis"], axis=1)
label_encoder = LabelEncoder()
y = label_encoder.fit_transform(df["diagnosis"])

X_train, X_test, y_train, y_test = train_test_split(X, y, random_state=42, test_size=0.2)

# RF model
rf_model = RandomForestClassifier()
rf_model.fit(X_train, y_train)
rf_y_hat = rf_model.predict(X_test)

rf_acc= accuracy_score(y_test, rf_y_hat)
rf_report = classification_report(y_test, rf_y_hat, target_names=label_encoder.classes_, output_dict=True)
rf_conf_mat = confusion_matrix(y_test, rf_y_hat)
disp = ConfusionMatrixDisplay(confusion_matrix=rf_conf_mat, display_labels=label_encoder.classes_)
disp.plot(cmap=plt.cm.Blues)
plt.title("RF confusion matrix - SMOTETomek")
plt.xticks(rotation=40, fontsize=5)
plt.yticks(fontsize=5)
plt.show()
plt.clf()


doc.add_heading("RF reports:")
rf_clf_table = pd.DataFrame(rf_report).transpose()
table = doc.add_table(rows=rf_clf_table.shape[0]+1, cols=rf_clf_table.shape[1]+1)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class/Metric'
for i, col_name in enumerate(rf_clf_table.columns):
    hdr_cells[i+1].text = col_name

for i, (index, row) in enumerate(rf_clf_table.iterrows()):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = str(index)
    for j, value in enumerate(row):
        row_cells[j+1].text = f"{value:.2f}" if isinstance(value, (float, np.float64)) else str(value)

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="single" w:sz="6"/>'
                            r'<w:left w:val="single" w:sz="6"/>'
                            r'<w:bottom w:val="single" w:sz="6"/>'
                            r'<w:right w:val="single" w:sz="6"/>'
                            r'</w:tcBorders>')
        cell_xml.append(borders)

doc.add_heading("RF accuracy", level=1)
doc.add_paragraph(rf_acc.__str__())

doc.add_paragraph("--------------------------------------------------------------------------------------------")
# ADABoost model

ada_model = AdaBoostClassifier()
ada_model.fit(X_train, y_train)
ada_y_hat = ada_model.predict(X_test)


ada_acc= accuracy_score(y_test, ada_y_hat)
ada_report = classification_report(y_test, ada_y_hat, target_names=label_encoder.classes_, output_dict=True)
ada_conf_mat = confusion_matrix(y_test, ada_y_hat)
disp = ConfusionMatrixDisplay(confusion_matrix=ada_conf_mat, display_labels=label_encoder.classes_)
disp.plot(cmap=plt.cm.Blues)
plt.title("ADABoost confusion matrix - SMOTETomek")
plt.xticks(rotation=40, fontsize=5)
plt.yticks(fontsize=5)
plt.show()
plt.clf()


doc.add_heading("ADABoost reports:")
ada_clf_table = pd.DataFrame(ada_report).transpose()
table = doc.add_table(rows=ada_clf_table.shape[0]+1, cols=ada_clf_table.shape[1]+1)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class/Metric'
for i, col_name in enumerate(ada_clf_table.columns):
    hdr_cells[i+1].text = col_name

for i, (index, row) in enumerate(ada_clf_table.iterrows()):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = str(index)
    for j, value in enumerate(row):
        row_cells[j+1].text = f"{value:.2f}" if isinstance(value, (float, np.float64)) else str(value)

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="single" w:sz="6"/>'
                            r'<w:left w:val="single" w:sz="6"/>'
                            r'<w:bottom w:val="single" w:sz="6"/>'
                            r'<w:right w:val="single" w:sz="6"/>'
                            r'</w:tcBorders>')
        cell_xml.append(borders)

doc.add_heading("ADABoost accuracy", level=1)
doc.add_paragraph(ada_acc.__str__())

doc.add_paragraph("--------------------------------------------------------------------------------------------")
# xg model
xg_model = XGBClassifier()
xg_model.fit(X_train, y_train)
xg_y_hat = xg_model.predict(X_test)

xg_acc= accuracy_score(y_test, xg_y_hat)
xg_report = classification_report(y_test, xg_y_hat, target_names=label_encoder.classes_, output_dict=True)
xg_conf_mat = confusion_matrix(y_test, xg_y_hat)
disp = ConfusionMatrixDisplay(confusion_matrix=xg_conf_mat, display_labels=label_encoder.classes_)
disp.plot(cmap=plt.cm.Blues)
plt.title("XGBoost confusion matrix - SMOTETomek")
plt.xticks(rotation=40, fontsize=5)
plt.yticks(fontsize=5)
plt.show()
plt.clf()


doc.add_heading("XGBoost reports:")
xg_clf_table = pd.DataFrame(xg_report).transpose()
table = doc.add_table(rows=xg_clf_table.shape[0]+1, cols=xg_clf_table.shape[1]+1)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class/Metric'
for i, col_name in enumerate(xg_clf_table.columns):
    hdr_cells[i+1].text = col_name

for i, (index, row) in enumerate(xg_clf_table.iterrows()):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = str(index)
    for j, value in enumerate(row):
        row_cells[j+1].text = f"{value:.2f}" if isinstance(value, (float, np.float64)) else str(value)

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="single" w:sz="6"/>'
                            r'<w:left w:val="single" w:sz="6"/>'
                            r'<w:bottom w:val="single" w:sz="6"/>'
                            r'<w:right w:val="single" w:sz="6"/>'
                            r'</w:tcBorders>')
        cell_xml.append(borders)

doc.add_heading("XGBoost accuracy", level=1)
doc.add_paragraph(xg_acc.__str__())


doc.add_paragraph("--------------------------------------------------------------------------------------------")
# lgb model
lgb_model = LGBMClassifier()
lgb_model.fit(X_train, y_train)
lgb_y_hat = lgb_model.predict(X_test)

lgb_acc= accuracy_score(y_test, lgb_y_hat)
lgb_report = classification_report(y_test, lgb_y_hat, target_names=label_encoder.classes_, output_dict=True)
lgb_conf_mat = confusion_matrix(y_test, lgb_y_hat)
disp = ConfusionMatrixDisplay(confusion_matrix=lgb_conf_mat, display_labels=label_encoder.classes_)
disp.plot(cmap=plt.cm.Blues)
plt.title("LGBM confusion matrix - SMOTETomek")
plt.xticks(rotation=40, fontsize=5)
plt.yticks(fontsize=5)
plt.show()
plt.clf()


doc.add_heading("LGBM reports:")
lgb_clf_table = pd.DataFrame(lgb_report).transpose()
table = doc.add_table(rows=lgb_clf_table.shape[0]+1, cols=lgb_clf_table.shape[1]+1)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class/Metric'
for i, col_name in enumerate(lgb_clf_table.columns):
    hdr_cells[i+1].text = col_name

for i, (index, row) in enumerate(lgb_clf_table.iterrows()):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = str(index)
    for j, value in enumerate(row):
        row_cells[j+1].text = f"{value:.2f}" if isinstance(value, (float, np.float64)) else str(value)

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="single" w:sz="6"/>'
                            r'<w:left w:val="single" w:sz="6"/>'
                            r'<w:bottom w:val="single" w:sz="6"/>'
                            r'<w:right w:val="single" w:sz="6"/>'
                            r'</w:tcBorders>')
        cell_xml.append(borders)

doc.add_heading("LGBM accuracy", level=1)
doc.add_paragraph(lgb_acc.__str__())


doc.save(f"SMOTETomek-report/{title}.docx")
